"""
=================================================================
SIGPDA - Reportes de experimentación de IA (PDF, Word, Excel)
=================================================================
Genera reportes de una ejecución de comparación de 5 modelos ya
persistida en BD (ver ia.persistencia_bd / database.modelos.
EjecucionEntrenamiento). No reemplaza los reportes generales
existentes (reportes/generador_pdf.py, reportes/exportador.py):
es un módulo aditivo, específico de experimentación de IA.

Cada reporte incluye: portada, fecha de generación, plato, rango del
dataset, resumen del EDA, estadísticas descriptivas, configuración e
hiperparámetros de los modelos, resultados de validación cruzada
(folds), tabla comparativa de los 5 modelos, modelo ganador, pruebas
estadísticas (Friedman/Wilcoxon/Diebold-Mariano), predicciones
futuras e interpretación/conclusiones automáticas.
"""

from __future__ import annotations

import io
import json
import logging
import os
import tempfile
from datetime import datetime

import matplotlib
matplotlib.use("Agg")  # backend no interactivo: seguro en un proceso de servidor
import matplotlib.pyplot as plt

from database.conexion import obtener_sesion
from database.modelos import (
    EjecucionEntrenamiento,
    ResultadoModelo,
    ResultadoFold,
    ResultadoPruebaEstadistica,
    Plato,
)

NOMBRES_LEGIBLES = {
    "arima": "ARIMA",
    "prophet": "Prophet",
    "holt_winters": "Holt-Winters",
    "transformer_random_forest": "Transformer + Random Forest",
    "transformer_gradient_boosting": "Transformer + Gradient Boosting",
}

logger = logging.getLogger(__name__)

DIAS_SEMANA_ABREV = ["Lun", "Mar", "Mie", "Jue", "Vie", "Sab", "Dom"]


# ================================================================
# LECTURA DE DATOS DESDE BD
# ================================================================

def obtener_datos_ejecucion(id_ejecucion: int) -> dict | None:
    """
    Reconstruye toda la información de una ejecución de comparación
    de 5 modelos desde las tablas normalizadas (Fase 7), sin volver a
    entrenar nada. Retorna None si el id_ejecucion no existe.
    """
    sesion = obtener_sesion()
    try:
        ejecucion = sesion.get(EjecucionEntrenamiento, id_ejecucion)
        if ejecucion is None:
            return None

        plato = sesion.query(Plato).filter(Plato.id_plato == ejecucion.id_plato).first()

        resultados_modelo = (
            sesion.query(ResultadoModelo)
            .filter(ResultadoModelo.id_ejecucion == id_ejecucion)
            .all()
        )
        resultados_modelo.sort(key=lambda r: (r.posicion is None, r.posicion or 0))

        modelos = []
        for rm in resultados_modelo:
            folds = (
                sesion.query(ResultadoFold)
                .filter(ResultadoFold.id_resultado_modelo == rm.id_resultado_modelo)
                .order_by(ResultadoFold.numero_fold)
                .all()
            )
            try:
                hp = json.loads(rm.hiperparametros_json) if rm.hiperparametros_json else {}
            except (TypeError, ValueError):
                hp = {}

            modelos.append({
                "modelo": rm.modelo,
                "modelo_legible": NOMBRES_LEGIBLES.get(rm.modelo, rm.modelo),
                "categoria": rm.categoria,
                "mae": rm.mae, "rmse": rm.rmse, "mape": rm.mape, "smape": rm.smape, "r2": rm.r2,
                "posicion": rm.posicion,
                "error_mensaje": rm.error_mensaje,
                "hiperparametros": hp,
                "folds": [
                    {
                        "numero_fold": f.numero_fold,
                        "fecha_inicio_train": f.fecha_inicio_train.isoformat() if f.fecha_inicio_train else None,
                        "fecha_fin_train": f.fecha_fin_train.isoformat() if f.fecha_fin_train else None,
                        "fecha_inicio_val": f.fecha_inicio_val.isoformat() if f.fecha_inicio_val else None,
                        "fecha_fin_val": f.fecha_fin_val.isoformat() if f.fecha_fin_val else None,
                        "n_train": f.n_train, "n_val": f.n_val,
                        "mae": f.mae, "rmse": f.rmse, "mape": f.mape, "smape": f.smape, "r2": f.r2,
                        "tiempo_entrenamiento": f.tiempo_entrenamiento,
                        "error_mensaje": f.error_mensaje,
                    }
                    for f in folds
                ],
            })

        pruebas = (
            sesion.query(ResultadoPruebaEstadistica)
            .filter(ResultadoPruebaEstadistica.id_ejecucion == id_ejecucion)
            .all()
        )
        pruebas_estadisticas = [
            {
                "prueba": p.prueba, "modelo_a": p.modelo_a, "modelo_b": p.modelo_b,
                "estadistico": p.estadistico, "p_valor": p.p_valor,
                "p_valor_ajustado": p.p_valor_ajustado, "significativo": p.significativo,
                "interpretacion": p.interpretacion,
            }
            for p in pruebas
        ]

        try:
            metadata_extra = json.loads(ejecucion.metadata_json) if ejecucion.metadata_json else {}
        except (TypeError, ValueError):
            metadata_extra = {}

        return {
            "id_ejecucion": ejecucion.id_ejecucion,
            "id_plato": ejecucion.id_plato,
            "nombre_plato": plato.nombre if plato else str(ejecucion.id_plato),
            "fecha": ejecucion.fecha,
            "estado": ejecucion.estado,
            "modelo_ganador": ejecucion.modelo_ganador,
            "modelo_ganador_legible": metadata_extra.get(
                "modelo_ganador_legible", NOMBRES_LEGIBLES.get(ejecucion.modelo_ganador, ejecucion.modelo_ganador),
            ),
            "mae_ganador": metadata_extra.get("mae_ganador"),
            "criterio_seleccion": ejecucion.criterio_seleccion,
            "duracion_segundos": ejecucion.duracion_segundos,
            "numero_registros": ejecucion.numero_registros,
            "fecha_inicio_datos": ejecucion.fecha_inicio_datos,
            "fecha_fin_datos": ejecucion.fecha_fin_datos,
            "dias_adelante": ejecucion.dias_adelante,
            "clima": ejecucion.clima,
            "evento": ejecucion.evento,
            "modelos": modelos,
            "pruebas_estadisticas": pruebas_estadisticas,
            "eda_resumen": metadata_extra.get("eda_resumen") or {},
            "predicciones_futuras": metadata_extra.get("predicciones_futuras") or [],
            "interpretacion": metadata_extra.get("interpretacion") or [],
            "hiperparametros_tuning": metadata_extra.get("hiperparametros_tuning") or {},
        }
    finally:
        sesion.close()


def _fmt(v, dec=4) -> str:
    if v is None:
        return "-"
    try:
        return f"{float(v):.{dec}f}"
    except (TypeError, ValueError):
        return str(v)


# ================================================================
# GENERACIÓN DE GRÁFICOS (matplotlib -> PNG temporal)
# ================================================================
# Se usan tanto en el PDF (fpdf2 embebe imágenes) como en Word
# (python-docx embebe imágenes). Cada función retorna la ruta de un
# archivo PNG temporal o None si no hay datos suficientes para
# graficar; el llamador es responsable de borrar el archivo con
# `_borrar_temporales` una vez insertado en el documento.

_COLOR_MORADO = "#6366F1"
_COLOR_VERDE = "#16A34A"
_COLOR_GRIS = "#9CA3AF"


def _nuevo_png_temporal() -> str:
    fd, ruta = tempfile.mkstemp(suffix=".png", prefix="sigpda_ia_")
    os.close(fd)
    return ruta


def _borrar_temporales(rutas: list[str]) -> None:
    for ruta in rutas:
        try:
            if ruta and os.path.exists(ruta):
                os.remove(ruta)
        except OSError as exc:
            logger.warning(f"No se pudo eliminar el gráfico temporal {ruta}: {exc}")


def _grafico_comparacion_mae(modelos: list[dict], modelo_ganador: str) -> str | None:
    """Barras horizontales con el MAE de cada modelo; resalta al ganador."""
    datos = [m for m in modelos if m.get("mae") is not None]
    if not datos:
        return None
    nombres = [m["modelo_legible"] for m in datos]
    valores = [float(m["mae"]) for m in datos]
    colores = [_COLOR_VERDE if m["modelo"] == modelo_ganador else _COLOR_MORADO for m in datos]

    fig, ax = plt.subplots(figsize=(7.2, 0.55 * len(datos) + 1), dpi=130)
    barras = ax.barh(nombres, valores, color=colores)
    ax.set_xlabel("MAE (menor es mejor)")
    ax.set_title("Comparación de modelos por MAE")
    ax.invert_yaxis()
    ax.bar_label(barras, fmt="%.3f", padding=3, fontsize=8)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()

    ruta = _nuevo_png_temporal()
    fig.savefig(ruta)
    plt.close(fig)
    return ruta


def _grafico_correlacion(correlaciones: dict) -> str | None:
    """Mapa de calor de la matriz de correlación entre variables numéricas."""
    if not correlaciones:
        return None
    columnas = list(correlaciones.keys())
    if len(columnas) < 2:
        return None
    matriz = [[correlaciones[c].get(c2) for c2 in columnas] for c in columnas]
    matriz = [[v if v is not None else 0.0 for v in fila] for fila in matriz]

    fig, ax = plt.subplots(figsize=(1.1 * len(columnas) + 2, 1.1 * len(columnas) + 1), dpi=130)
    im = ax.imshow(matriz, cmap="RdBu_r", vmin=-1, vmax=1)
    ax.set_xticks(range(len(columnas)))
    ax.set_yticks(range(len(columnas)))
    ax.set_xticklabels(columnas, rotation=45, ha="right", fontsize=8)
    ax.set_yticklabels(columnas, fontsize=8)
    for i in range(len(columnas)):
        for j in range(len(columnas)):
            valor = matriz[i][j]
            color_texto = "white" if abs(valor) > 0.6 else "black"
            ax.text(j, i, f"{valor:.2f}", ha="center", va="center", color=color_texto, fontsize=7.5)
    ax.set_title("Mapa de calor de correlaciones")
    fig.colorbar(im, ax=ax, shrink=0.8, label="Coeficiente de correlación")
    fig.tight_layout()

    ruta = _nuevo_png_temporal()
    fig.savefig(ruta)
    plt.close(fig)
    return ruta


def _grafico_serie_historica(serie_historica: list[dict]) -> str | None:
    """Evolución histórica de la demanda (fecha vs. cantidad)."""
    if not serie_historica:
        return None
    fechas = [p["fecha"] for p in serie_historica]
    valores = [p["cantidad"] for p in serie_historica]

    fig, ax = plt.subplots(figsize=(7.2, 2.8), dpi=130)
    ax.plot(fechas, valores, color=_COLOR_MORADO, linewidth=1)
    ax.set_title("Evolución histórica de la demanda")
    ax.set_ylabel("Cantidad vendida")
    n = len(fechas)
    paso = max(1, n // 10)
    ax.set_xticks(range(0, n, paso))
    ax.set_xticklabels([fechas[i] for i in range(0, n, paso)], rotation=45, ha="right", fontsize=7)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()

    ruta = _nuevo_png_temporal()
    fig.savefig(ruta)
    plt.close(fig)
    return ruta


def _grafico_dia_semana(por_dia_semana: list[dict]) -> str | None:
    """Demanda promedio por día de la semana."""
    if not por_dia_semana:
        return None
    dias = [d["dia_semana"] for d in por_dia_semana]
    valores = [d["demanda_promedio"] for d in por_dia_semana]

    fig, ax = plt.subplots(figsize=(7.2, 2.6), dpi=130)
    ax.bar(dias, valores, color=_COLOR_MORADO)
    ax.set_title("Demanda promedio por día de la semana")
    ax.set_ylabel("Demanda promedio")
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()

    ruta = _nuevo_png_temporal()
    fig.savefig(ruta)
    plt.close(fig)
    return ruta


# ================================================================
# PDF (fpdf2)
# ================================================================

def _n(txt) -> str:
    """Normaliza caracteres latinos para fpdf2 (Latin-1), igual que reportes/reportes.py."""
    txt = str(txt)
    return (txt
        .replace("á", "a").replace("é", "e").replace("í", "i")
        .replace("ó", "o").replace("ú", "u").replace("ü", "u")
        .replace("ñ", "n").replace("Á", "A").replace("É", "E")
        .replace("Í", "I").replace("Ó", "O").replace("Ú", "U")
        .replace("Ñ", "N").replace("°", "").replace("–", "-")
        .replace("’", "'").replace("“", '"').replace("”", '"')
        .replace("²", "2")
    )


def _construir_pdf(datos: dict, temporales: list[str]):
    from fpdf import FPDF

    _MORADO = (99, 102, 241)
    _GRIS = (107, 114, 128)
    _GRIS_F = (243, 244, 246)
    _VERDE = (22, 163, 74)

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(10, 15, 10)

    def seccion(titulo):
        pdf.ln(3)
        pdf.set_fill_color(*_MORADO)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 8, f"  {_n(titulo)}", fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(2)

    def parrafo(texto, size=9, color=(0, 0, 0)):
        pdf.set_font("Helvetica", "", size)
        pdf.set_text_color(*color)
        pdf.multi_cell(0, 5, _n(texto), new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)

    def tabla(encabezados, filas, anchos):
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_fill_color(*_MORADO)
        pdf.set_text_color(255, 255, 255)
        for h, w in zip(encabezados, anchos):
            pdf.cell(w, 6, _n(h), border=1, fill=True, align="C")
        pdf.ln()
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 7.5)
        for i, fila in enumerate(filas):
            pdf.set_fill_color(*_GRIS_F) if i % 2 == 0 else pdf.set_fill_color(255, 255, 255)
            for celda, w in zip(fila, anchos):
                pdf.cell(w, 5.5, _n(celda), border=1, fill=True)
            pdf.ln()
        pdf.ln(1)

    def imagen(ruta, ancho=180, alto_estimado=75):
        if not ruta:
            return
        if pdf.get_y() + alto_estimado > 275:
            pdf.add_page()
        pdf.image(ruta, x=(210 - ancho) / 2, w=ancho)
        pdf.ln(3)

    # --- Portada / encabezado ---
    pdf.set_fill_color(*_MORADO)
    pdf.rect(0, 0, 210, 20, "F")
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(255, 255, 255)
    pdf.set_y(5)
    pdf.cell(0, 10, "SIGPDA - Reporte de Experimentacion de IA", align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.set_y(24)

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*_GRIS)
    pdf.cell(0, 6, f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(2)

    # --- Resumen ---
    seccion("Resumen de la ejecucion")
    parrafo(
        f"Plato analizado: {datos['nombre_plato']} (ID {datos['id_plato']})  |  "
        f"Ejecucion #{datos['id_ejecucion']}  |  Estado: {datos['estado']}"
    )
    parrafo(
        f"Rango del dataset: {datos['fecha_inicio_datos']} a {datos['fecha_fin_datos']}  |  "
        f"Registros: {datos['numero_registros']}  |  Duracion: {_fmt(datos['duracion_segundos'], 2)} s"
    )
    parrafo(
        f"Modelo ganador: {datos['modelo_ganador_legible']}  (MAE = {_fmt(datos['mae_ganador'])})  |  "
        f"Criterio: {datos['criterio_seleccion']}",
        color=_VERDE,
    )

    # --- EDA resumido ---
    eda = datos.get("eda_resumen") or {}
    if eda.get("resumen") or eda.get("estadisticas_descriptivas"):
        seccion("Resumen del analisis exploratorio de datos (EDA)")
        stats = eda.get("estadisticas_descriptivas", {})
        if stats:
            parrafo(
                f"Media: {_fmt(stats.get('media'))}  Mediana: {_fmt(stats.get('mediana'))}  "
                f"Desv. estandar: {_fmt(stats.get('desviacion_estandar'))}  "
                f"Minimo: {_fmt(stats.get('minimo'))}  Maximo: {_fmt(stats.get('maximo'))}  "
                f"Q1: {_fmt(stats.get('q1'))}  Q3: {_fmt(stats.get('q3'))}"
            )
        for adv in eda.get("advertencias", []):
            parrafo(f"* {adv}", size=8, color=_GRIS)

        ruta_serie = _grafico_serie_historica(eda.get("serie_historica") or [])
        ruta_dia = _grafico_dia_semana(eda.get("por_dia_semana") or [])
        ruta_corr = _grafico_correlacion(eda.get("correlaciones") or {})
        for r in (ruta_serie, ruta_dia, ruta_corr):
            if r:
                temporales.append(r)
        imagen(ruta_serie, ancho=180, alto_estimado=65)
        imagen(ruta_dia, ancho=170, alto_estimado=60)
        imagen(ruta_corr, ancho=140, alto_estimado=110)

    # --- Configuracion e hiperparametros ---
    seccion("Configuracion e hiperparametros por modelo")
    filas_hp = []
    for m in datos["modelos"]:
        hp_texto = ", ".join(f"{k}={v}" for k, v in (m["hiperparametros"].get("hiperparametros") or m["hiperparametros"]).items()) \
            if isinstance(m["hiperparametros"], dict) else str(m["hiperparametros"])
        filas_hp.append([m["modelo_legible"], (hp_texto or "-")[:70]])
    tabla(["Modelo", "Configuracion utilizada"], filas_hp, [55, 135])

    tuning = datos.get("hiperparametros_tuning") or {}
    if tuning:
        seccion("Optimizacion de hiperparametros (mejor combinacion por modelo)")
        filas_tuning = []
        for nombre, t in tuning.items():
            if not isinstance(t, dict) or not t.get("aplicable"):
                continue
            mejor = t.get("mejor_hiperparametros") or {}
            mejor_txt = ", ".join(f"{k}={v}" for k, v in mejor.items())
            filas_tuning.append([
                NOMBRES_LEGIBLES.get(nombre, nombre), (mejor_txt or "-")[:55],
                _fmt(t.get("mejor_valor")), str(t.get("n_combinaciones", "-")),
            ])
        if filas_tuning:
            tabla(["Modelo", "Mejor configuracion", "Mejor MAE", "Combin."], filas_tuning, [45, 100, 25, 20])

    # --- Comparacion de los 5 modelos ---
    seccion("Comparacion de modelos")
    filas_cmp = []
    for m in datos["modelos"]:
        marca = " (GANADOR)" if m["modelo"] == datos["modelo_ganador"] else ""
        filas_cmp.append([
            m["modelo_legible"] + marca, _fmt(m["mae"]), _fmt(m["rmse"]),
            _fmt(m["mape"], 2), _fmt(m["smape"], 2), _fmt(m["r2"]),
        ])
    tabla(["Modelo", "MAE", "RMSE", "MAPE %", "SMAPE %", "R2"], filas_cmp, [60, 26, 26, 26, 26, 26])

    ruta_mae = _grafico_comparacion_mae(datos["modelos"], datos["modelo_ganador"])
    if ruta_mae:
        temporales.append(ruta_mae)
    imagen(ruta_mae, ancho=170, alto_estimado=70)

    # --- Validacion cruzada (folds) ---
    seccion("Validacion cruzada temporal (folds)")
    for m in datos["modelos"]:
        if not m["folds"]:
            continue
        parrafo(f"{m['modelo_legible']}:", size=8.5)
        filas_folds = [
            [
                str(f["numero_fold"]), f"{f['n_train']}/{f['n_val']}",
                _fmt(f["mae"], 3), _fmt(f["rmse"], 3), _fmt(f["smape"], 2),
            ]
            for f in m["folds"]
        ]
        tabla(["Fold", "N train/val", "MAE", "RMSE", "SMAPE %"], filas_folds, [20, 35, 35, 35, 39])

    # --- Pruebas estadisticas ---
    seccion("Pruebas estadisticas (alpha = 0.05)")
    for p in datos["pruebas_estadisticas"]:
        sig = "SIGNIFICATIVO" if p["significativo"] else "no significativo"
        etiqueta = f"{p['prueba'].upper()}"
        if p.get("modelo_a") and p.get("modelo_b"):
            etiqueta += f" ({p['modelo_a']} vs {p['modelo_b']})"
        parrafo(
            f"{etiqueta}: estadistico={_fmt(p['estadistico'])}  p-valor={_fmt(p['p_valor'])}  "
            f"p-ajustado={_fmt(p['p_valor_ajustado']) if p['p_valor_ajustado'] is not None else '-'}  [{sig}]",
            size=8,
        )
        if p.get("interpretacion"):
            parrafo(f"  {p['interpretacion']}", size=7.5, color=_GRIS)

    # --- Predicciones futuras ---
    if datos["predicciones_futuras"]:
        seccion("Predicciones futuras (modelo ganador)")
        filas_pred = [
            [p["fecha"], str(p["demanda_estimada"]), str(p["recomendacion"]), p["riesgo"]]
            for p in datos["predicciones_futuras"]
        ]
        tabla(["Fecha", "Demanda est.", "Prod. recomendada", "Riesgo"], filas_pred, [45, 45, 55, 45])

    # --- Interpretacion / conclusiones ---
    if datos["interpretacion"]:
        seccion("Interpretacion y conclusiones")
        for texto in datos["interpretacion"]:
            parrafo(f"- {texto}", size=8.5)

    return pdf


def generar_pdf_ia(id_ejecucion: int) -> bytes | None:
    """Genera el reporte PDF de una ejecución de comparación de 5 modelos."""
    datos = obtener_datos_ejecucion(id_ejecucion)
    if datos is None:
        return None
    temporales: list[str] = []
    try:
        pdf = _construir_pdf(datos, temporales)
        return bytes(pdf.output())
    finally:
        _borrar_temporales(temporales)


# ================================================================
# WORD (python-docx)
# ================================================================

def generar_word_ia(id_ejecucion: int) -> bytes | None:
    """Genera el reporte Word de una ejecución de comparación de 5 modelos."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    datos = obtener_datos_ejecucion(id_ejecucion)
    if datos is None:
        return None

    temporales: list[str] = []

    def imagen(ruta, ancho_pulgadas=6.0):
        if ruta:
            doc.add_picture(ruta, width=Inches(ancho_pulgadas))

    doc = Document()

    titulo = doc.add_heading("SIGPDA — Reporte de Experimentación de IA", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fecha = doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Resumen de la ejecución", level=2)
    doc.add_paragraph(f"Plato analizado: {datos['nombre_plato']} (ID {datos['id_plato']})")
    doc.add_paragraph(f"Ejecución #{datos['id_ejecucion']} — Estado: {datos['estado']}")
    doc.add_paragraph(
        f"Rango del dataset: {datos['fecha_inicio_datos']} a {datos['fecha_fin_datos']} "
        f"({datos['numero_registros']} registros)"
    )
    doc.add_paragraph(f"Duración total: {_fmt(datos['duracion_segundos'], 2)} s")
    p_ganador = doc.add_paragraph()
    p_ganador.add_run(
        f"Modelo ganador: {datos['modelo_ganador_legible']} (MAE = {_fmt(datos['mae_ganador'])})"
    ).bold = True
    doc.add_paragraph(f"Criterio de selección: {datos['criterio_seleccion']}")

    # --- EDA ---
    eda = datos.get("eda_resumen") or {}
    stats = eda.get("estadisticas_descriptivas") or {}
    if stats:
        doc.add_heading("Resumen del análisis exploratorio de datos (EDA)", level=2)
        tabla_eda = doc.add_table(rows=1, cols=2)
        tabla_eda.style = "Light Grid Accent 1"
        tabla_eda.rows[0].cells[0].text = "Estadística"
        tabla_eda.rows[0].cells[1].text = "Valor"
        for k, v in stats.items():
            fila = tabla_eda.add_row().cells
            fila[0].text = str(k)
            fila[1].text = _fmt(v)
        for adv in eda.get("advertencias", []):
            doc.add_paragraph(adv, style="Intense Quote")

        ruta_serie = _grafico_serie_historica(eda.get("serie_historica") or [])
        ruta_dia = _grafico_dia_semana(eda.get("por_dia_semana") or [])
        ruta_corr = _grafico_correlacion(eda.get("correlaciones") or {})
        for r in (ruta_serie, ruta_dia, ruta_corr):
            if r:
                temporales.append(r)
        imagen(ruta_serie)
        imagen(ruta_dia)
        imagen(ruta_corr, ancho_pulgadas=5.0)

    # --- Comparación de modelos ---
    doc.add_heading("Comparación de modelos", level=2)
    tabla_cmp = doc.add_table(rows=1, cols=6)
    tabla_cmp.style = "Light Grid Accent 1"
    for i, h in enumerate(["Modelo", "MAE", "RMSE", "MAPE %", "SMAPE %", "R²"]):
        tabla_cmp.rows[0].cells[i].text = h
    for m in datos["modelos"]:
        fila = tabla_cmp.add_row().cells
        marca = " (GANADOR)" if m["modelo"] == datos["modelo_ganador"] else ""
        fila[0].text = m["modelo_legible"] + marca
        fila[1].text = _fmt(m["mae"])
        fila[2].text = _fmt(m["rmse"])
        fila[3].text = _fmt(m["mape"], 2)
        fila[4].text = _fmt(m["smape"], 2)
        fila[5].text = _fmt(m["r2"])

    ruta_mae = _grafico_comparacion_mae(datos["modelos"], datos["modelo_ganador"])
    if ruta_mae:
        temporales.append(ruta_mae)
    imagen(ruta_mae, ancho_pulgadas=5.5)

    # --- Hiperparámetros / tuning ---
    tuning = datos.get("hiperparametros_tuning") or {}
    aplicables = {k: v for k, v in tuning.items() if isinstance(v, dict) and v.get("aplicable")}
    if aplicables:
        doc.add_heading("Optimización de hiperparámetros", level=2)
        for nombre, t in aplicables.items():
            doc.add_paragraph(
                f"{NOMBRES_LEGIBLES.get(nombre, nombre)}: mejor MAE = {_fmt(t.get('mejor_valor'))}, "
                f"{t.get('n_combinaciones', 0)} combinaciones evaluadas, semilla = {t.get('semilla')}"
            )
            mejor = t.get("mejor_hiperparametros") or {}
            if mejor:
                doc.add_paragraph(
                    "Mejor configuración: " + ", ".join(f"{k}={v}" for k, v in mejor.items()),
                    style="List Bullet",
                )

    # --- Validación cruzada ---
    doc.add_heading("Validación cruzada temporal (folds)", level=2)
    for m in datos["modelos"]:
        if not m["folds"]:
            continue
        doc.add_heading(m["modelo_legible"], level=3)
        tabla_folds = doc.add_table(rows=1, cols=5)
        tabla_folds.style = "Light Grid Accent 1"
        for i, h in enumerate(["Fold", "N train/val", "MAE", "RMSE", "SMAPE %"]):
            tabla_folds.rows[0].cells[i].text = h
        for f in m["folds"]:
            fila = tabla_folds.add_row().cells
            fila[0].text = str(f["numero_fold"])
            fila[1].text = f"{f['n_train']}/{f['n_val']}"
            fila[2].text = _fmt(f["mae"], 3)
            fila[3].text = _fmt(f["rmse"], 3)
            fila[4].text = _fmt(f["smape"], 2)

    # --- Pruebas estadísticas ---
    doc.add_heading("Pruebas estadísticas (alpha = 0.05)", level=2)
    for p in datos["pruebas_estadisticas"]:
        sig = "significativo" if p["significativo"] else "no significativo"
        etiqueta = p["prueba"].capitalize()
        if p.get("modelo_a") and p.get("modelo_b"):
            etiqueta += f" ({p['modelo_a']} vs {p['modelo_b']})"
        doc.add_paragraph(
            f"{etiqueta}: estadístico={_fmt(p['estadistico'])}, p-valor={_fmt(p['p_valor'])}, "
            f"p-ajustado={_fmt(p['p_valor_ajustado']) if p['p_valor_ajustado'] is not None else '—'} [{sig}]"
        )
        if p.get("interpretacion"):
            doc.add_paragraph(p["interpretacion"], style="Intense Quote")

    # --- Predicciones futuras ---
    if datos["predicciones_futuras"]:
        doc.add_heading("Predicciones futuras", level=2)
        tabla_pred = doc.add_table(rows=1, cols=4)
        tabla_pred.style = "Light Grid Accent 1"
        for i, h in enumerate(["Fecha", "Demanda est.", "Prod. recomendada", "Riesgo"]):
            tabla_pred.rows[0].cells[i].text = h
        for p in datos["predicciones_futuras"]:
            fila = tabla_pred.add_row().cells
            fila[0].text = str(p["fecha"])
            fila[1].text = str(p["demanda_estimada"])
            fila[2].text = str(p["recomendacion"])
            fila[3].text = str(p["riesgo"])

    # --- Interpretación ---
    if datos["interpretacion"]:
        doc.add_heading("Interpretación y conclusiones", level=2)
        for texto in datos["interpretacion"]:
            doc.add_paragraph(texto, style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    _borrar_temporales(temporales)
    return buffer.getvalue()


# ================================================================
# EXCEL (openpyxl, multi-hoja)
# ================================================================

def generar_excel_ia(id_ejecucion: int) -> bytes | None:
    """Genera el reporte Excel (multi-hoja) de una ejecución de comparación de 5 modelos."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    from openpyxl.chart import BarChart, LineChart, Reference
    from openpyxl.formatting.rule import ColorScaleRule

    datos = obtener_datos_ejecucion(id_ejecucion)
    if datos is None:
        return None

    wb = Workbook()

    fuente_encabezado = Font(bold=True, color="FFFFFF")
    relleno_encabezado = PatternFill("solid", fgColor="6366F1")

    def preparar_hoja(ws, encabezados, anchos):
        for col_idx, (h, ancho) in enumerate(zip(encabezados, anchos), start=1):
            celda = ws.cell(row=1, column=col_idx, value=h)
            celda.font = fuente_encabezado
            celda.fill = relleno_encabezado
            celda.alignment = Alignment(horizontal="center")
            ws.column_dimensions[get_column_letter(col_idx)].width = ancho
        ws.freeze_panes = "A2"

    # --- Hoja: Resumen ---
    ws = wb.active
    ws.title = "Resumen"
    ws.append(["Campo", "Valor"])
    for c in ws[1]:
        c.font = fuente_encabezado
        c.fill = relleno_encabezado
    for campo, valor in [
        ("Plato", datos["nombre_plato"]),
        ("ID plato", datos["id_plato"]),
        ("Ejecución", datos["id_ejecucion"]),
        ("Estado", datos["estado"]),
        ("Fecha inicio datos", str(datos["fecha_inicio_datos"])),
        ("Fecha fin datos", str(datos["fecha_fin_datos"])),
        ("Registros", datos["numero_registros"]),
        ("Duración (s)", datos["duracion_segundos"]),
        ("Modelo ganador", datos["modelo_ganador_legible"]),
        ("MAE ganador", datos["mae_ganador"]),
        ("Criterio de selección", datos["criterio_seleccion"]),
        ("Generado", datetime.now().strftime("%d/%m/%Y %H:%M")),
    ]:
        ws.append([campo, valor])
    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 40
    ws.freeze_panes = "A2"

    # --- Hoja: EDA ---
    ws_eda = wb.create_sheet("EDA")
    eda = datos.get("eda_resumen") or {}
    preparar_hoja(ws_eda, ["Estadística", "Valor"], [22, 16])
    for k, v in (eda.get("estadisticas_descriptivas") or {}).items():
        ws_eda.append([k, v])
    if eda.get("advertencias"):
        ws_eda.append([])
        ws_eda.append(["Advertencias"])
        for a in eda["advertencias"]:
            ws_eda.append([a])

    # --- Demanda por día de semana (tabla + gráfico) ---
    por_dia_semana = eda.get("por_dia_semana") or []
    if por_dia_semana:
        fila_inicio_dia = ws_eda.max_row + 3
        ws_eda.cell(row=fila_inicio_dia - 1, column=1, value="Demanda por día de la semana").font = Font(bold=True)
        ws_eda.cell(row=fila_inicio_dia, column=1, value="Día").font = fuente_encabezado
        ws_eda.cell(row=fila_inicio_dia, column=1).fill = relleno_encabezado
        ws_eda.cell(row=fila_inicio_dia, column=2, value="Demanda promedio").font = fuente_encabezado
        ws_eda.cell(row=fila_inicio_dia, column=2).fill = relleno_encabezado
        for i, d in enumerate(por_dia_semana, start=1):
            ws_eda.cell(row=fila_inicio_dia + i, column=1, value=d["dia_semana"])
            ws_eda.cell(row=fila_inicio_dia + i, column=2, value=d["demanda_promedio"])

        chart_dia = BarChart()
        chart_dia.title = "Demanda promedio por día de la semana"
        chart_dia.y_axis.title = "Demanda promedio"
        n_dia = len(por_dia_semana)
        datos_dia_ref = Reference(ws_eda, min_col=2, min_row=fila_inicio_dia, max_row=fila_inicio_dia + n_dia)
        cat_dia_ref = Reference(ws_eda, min_col=1, min_row=fila_inicio_dia + 1, max_row=fila_inicio_dia + n_dia)
        chart_dia.add_data(datos_dia_ref, titles_from_data=True)
        chart_dia.set_categories(cat_dia_ref)
        ws_eda.add_chart(chart_dia, f"D{fila_inicio_dia}")

    # --- Serie histórica (tabla + gráfico de línea) ---
    ws_serie = wb.create_sheet("Serie historica")
    serie_historica = eda.get("serie_historica") or []
    preparar_hoja(ws_serie, ["Fecha", "Cantidad", "Interpolado"], [14, 12, 12])
    for p in serie_historica:
        ws_serie.append([p["fecha"], p["cantidad"], "Si" if p.get("interpolado") else "No"])
    if serie_historica:
        chart_serie = LineChart()
        chart_serie.title = "Evolución histórica de la demanda"
        chart_serie.y_axis.title = "Cantidad"
        n_serie = len(serie_historica)
        datos_serie_ref = Reference(ws_serie, min_col=2, min_row=1, max_row=n_serie + 1)
        cat_serie_ref = Reference(ws_serie, min_col=1, min_row=2, max_row=n_serie + 1)
        chart_serie.add_data(datos_serie_ref, titles_from_data=True)
        chart_serie.set_categories(cat_serie_ref)
        chart_serie.width = 24
        chart_serie.height = 9
        ws_serie.add_chart(chart_serie, "E2")

    # --- Correlaciones (matriz con formato de mapa de calor) ---
    correlaciones = eda.get("correlaciones") or {}
    if correlaciones:
        ws_corr = wb.create_sheet("Correlaciones")
        columnas_corr = list(correlaciones.keys())
        ws_corr.cell(row=1, column=1, value="")
        for j, col in enumerate(columnas_corr, start=2):
            c = ws_corr.cell(row=1, column=j, value=col)
            c.font = fuente_encabezado
            c.fill = relleno_encabezado
        for i, fila_nombre in enumerate(columnas_corr, start=2):
            c = ws_corr.cell(row=i, column=1, value=fila_nombre)
            c.font = fuente_encabezado
            c.fill = relleno_encabezado
            for j, col in enumerate(columnas_corr, start=2):
                valor = correlaciones.get(fila_nombre, {}).get(col)
                celda = ws_corr.cell(row=i, column=j, value=valor if valor is not None else 0.0)
                celda.number_format = "0.00"
        n_corr = len(columnas_corr)
        rango_corr = f"B2:{get_column_letter(n_corr + 1)}{n_corr + 1}"
        ws_corr.conditional_formatting.add(
            rango_corr,
            ColorScaleRule(
                start_type="num", start_value=-1, start_color="EF4444",
                mid_type="num", mid_value=0, mid_color="FFFFFF",
                end_type="num", end_value=1, end_color="16A34A",
            ),
        )
        ws_corr.column_dimensions["A"].width = 20
        for j in range(2, n_corr + 2):
            ws_corr.column_dimensions[get_column_letter(j)].width = 14

    # --- Hoja: Modelos (comparación + gráfico) ---
    ws_mod = wb.create_sheet("Modelos")
    preparar_hoja(
        ws_mod,
        ["Modelo", "Categoría", "MAE", "RMSE", "MAPE %", "SMAPE %", "R2", "Posición", "Ganador"],
        [32, 12, 10, 10, 10, 10, 10, 10, 10],
    )
    for m in datos["modelos"]:
        ws_mod.append([
            m["modelo_legible"], m["categoria"], m["mae"], m["rmse"], m["mape"], m["smape"], m["r2"],
            m["posicion"], "Si" if m["modelo"] == datos["modelo_ganador"] else "",
        ])
    for fila in ws_mod.iter_rows(min_row=2, min_col=3, max_col=7):
        for celda in fila:
            celda.number_format = "0.0000"

    if datos["modelos"]:
        chart = BarChart()
        chart.title = "MAE por modelo"
        chart.y_axis.title = "MAE"
        n = len(datos["modelos"])
        datos_ref = Reference(ws_mod, min_col=3, min_row=1, max_row=n + 1)
        categorias_ref = Reference(ws_mod, min_col=1, min_row=2, max_row=n + 1)
        chart.add_data(datos_ref, titles_from_data=True)
        chart.set_categories(categorias_ref)
        ws_mod.add_chart(chart, "K2")

    # --- Hoja: Validación cruzada ---
    ws_cv = wb.create_sheet("Validacion cruzada")
    preparar_hoja(
        ws_cv,
        ["Modelo", "Fold", "N train", "N val", "MAE", "RMSE", "MAPE %", "SMAPE %", "R2", "Tiempo (s)"],
        [30, 8, 10, 10, 10, 10, 10, 10, 10, 12],
    )
    for m in datos["modelos"]:
        for f in m["folds"]:
            ws_cv.append([
                m["modelo_legible"], f["numero_fold"], f["n_train"], f["n_val"],
                f["mae"], f["rmse"], f["mape"], f["smape"], f["r2"], f["tiempo_entrenamiento"],
            ])

    # --- Hoja: Hiperparámetros ---
    ws_hp = wb.create_sheet("Hiperparametros")
    preparar_hoja(
        ws_hp,
        ["Modelo", "Mejor configuración", "Mejor MAE", "Combinaciones", "Tiempo (s)", "Semilla"],
        [32, 55, 12, 14, 12, 10],
    )
    for nombre, t in (datos.get("hiperparametros_tuning") or {}).items():
        if not isinstance(t, dict) or not t.get("aplicable"):
            continue
        mejor = t.get("mejor_hiperparametros") or {}
        ws_hp.append([
            NOMBRES_LEGIBLES.get(nombre, nombre),
            ", ".join(f"{k}={v}" for k, v in mejor.items()),
            t.get("mejor_valor"), t.get("n_combinaciones"), t.get("tiempo_total"), t.get("semilla"),
        ])

    # --- Hoja: Pruebas estadísticas ---
    ws_pe = wb.create_sheet("Pruebas estadisticas")
    preparar_hoja(
        ws_pe,
        ["Prueba", "Modelo A", "Modelo B", "Estadístico", "p-valor", "p-ajustado", "Significativo", "Interpretación"],
        [14, 22, 22, 12, 10, 12, 12, 70],
    )
    for p in datos["pruebas_estadisticas"]:
        ws_pe.append([
            p["prueba"], p.get("modelo_a"), p.get("modelo_b"), p["estadistico"], p["p_valor"],
            p["p_valor_ajustado"], "Si" if p["significativo"] else "No", p.get("interpretacion"),
        ])

    # --- Hoja: Predicciones ---
    ws_pred = wb.create_sheet("Predicciones")
    preparar_hoja(ws_pred, ["Fecha", "Demanda estimada", "Producción recomendada", "Riesgo"], [14, 18, 24, 12])
    for p in datos["predicciones_futuras"]:
        ws_pred.append([p["fecha"], p["demanda_estimada"], p["recomendacion"], p["riesgo"]])

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()
